# pyright: reportAttributeAccessIssue=false
"""Word/DOCX parsing mixin for AttachmentParser.

Handles Word document text extraction:
- .docx via python-docx
- .doc via antiword/catdoc/olefile fallback chain
"""

from __future__ import annotations

from pathlib import Path

from .models import AttachmentParseResult
from ._attachment_helpers import (
    _decode_ole_text,
    _text_chars,
    _try_cli_doc_extract,
)


class _WordParserMixin:
    """Word document parsing methods for AttachmentParser."""

    @staticmethod
    def _extract_doc_olefile(file_path: Path) -> str | None:
        """Pure Python .doc 텍스트 추출 (olefile OLE2 스트림 파싱)"""
        try:
            import olefile
        except ImportError:
            return None

        try:
            ole = olefile.OleFileIO(str(file_path))
        except (
            RuntimeError, OSError, ValueError,
            TypeError, KeyError, AttributeError,
        ):
            return None

        try:
            if not ole.exists("WordDocument"):
                return None

            raw = ole.openstream("WordDocument").read()
            if len(raw) < 20:
                return None

            return _decode_ole_text(raw)
        finally:
            ole.close()

    @staticmethod
    def _parse_legacy_doc(file_path: Path) -> AttachmentParseResult:
        """레거시 .doc (OLE2) 파일에서 텍스트 추출

        Strategy: antiword -> catdoc -> olefile fallback
        """
        import shutil

        # 1차: antiword (테이블 구조 보존, Docker 환경)
        result = _try_cli_doc_extract(
            shutil.which("antiword"), file_path, confidence=0.7,
        )
        if result is not None:
            return result

        # 2차: catdoc (antiword 없을 때, Docker 환경)
        result = _try_cli_doc_extract(
            shutil.which("catdoc"), file_path,
            confidence=0.6, extra_args=["-w"],
        )
        if result is not None:
            return result

        # 3차: olefile pure Python 추출 (로컬 환경)
        text = _WordParserMixin._extract_doc_olefile(file_path)
        if text:
            return AttachmentParseResult(
                extracted_text=text,
                extracted_tables=[],
                confidence=0.5,
                native_text_chars=_text_chars(text),
            )

        return AttachmentParseResult(
            extracted_text=(
                "[.doc 파싱 실패: antiword/catdoc 미설치"
                " 및 olefile 추출 실패]"
            ),
            extracted_tables=[],
            confidence=0.0,
            ocr_skip_reason="parse_error",
        )

    @staticmethod
    def _extract_word_tables(doc) -> list[dict]:
        """Extract tables from a python-docx Document object."""
        tables = []
        for idx, table in enumerate(doc.tables, 1):
            rows_data = []
            for row in table.rows:
                row_values = [cell.text.strip() for cell in row.cells]
                rows_data.append(row_values)

            if rows_data:
                headers = rows_data[0]
                data_rows = (
                    rows_data[1:] if len(rows_data) > 1 else []
                )
                tables.append({
                    "table_index": idx,
                    "headers": headers,
                    "rows": [
                        dict(zip(headers, row))
                        for row in data_rows
                        if len(row) == len(headers)
                    ],
                })
        return tables

    @classmethod
    def parse_word(cls, file_path: Path) -> AttachmentParseResult:
        """Word에서 텍스트와 테이블 추출

        .docx: python-docx, .doc: antiword/catdoc
        """
        try:
            if str(file_path).lower().endswith(".doc"):
                return cls._parse_legacy_doc(file_path)
            from docx import Document

            doc = Document(file_path)
            text_parts = [
                p.text for p in doc.paragraphs if p.text.strip()
            ]
            tables = cls._extract_word_tables(doc)

            full_text = "\n\n".join(text_parts)
            confidence = 0.9 if full_text.strip() else 0.0

            return AttachmentParseResult(
                extracted_text=full_text,
                extracted_tables=tables,
                confidence=confidence,
                native_text_chars=(
                    _text_chars(full_text)
                ),
            )

        except (
            RuntimeError, OSError, ValueError,
            TypeError, KeyError, AttributeError,
        ) as e:
            return AttachmentParseResult(
                extracted_text=f"[Word 파싱 오류: {e}]",
                extracted_tables=[],
                confidence=0.0,
                ocr_skip_reason="parse_error",
            )
