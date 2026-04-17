"""Document parser -- shared utilities, data classes, and common parsers.

Extracted from document_parser.py for SRP.
"""

from __future__ import annotations

import io
import logging
import shutil
import subprocess
import tempfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from src.config.weights import weights as _w

_EXT_PPTX = ".pptx"

logger = logging.getLogger(__name__)


@dataclass
class ParseResult:
    """Enhanced document parse result with text, tables, and OCR data."""

    text: str = ""
    tables: list[list[list[str]]] = field(default_factory=list)
    ocr_text: str = ""  # text extracted via OCR from images/scanned pages
    images_processed: int = 0
    visual_analyses: list[dict[str, Any]] = field(default_factory=list)
    file_modified_at: str = ""  # ISO timestamp from file metadata (PDF modDate, PPTX modified)

    @property
    def full_text(self) -> str:
        """Combined text from all sources."""
        parts = [self.text]
        if self.ocr_text:
            parts.append(f"\n[OCR Extracted Text]\n{self.ocr_text}")
        for table in self.tables:
            parts.append(f"\n[Table]\n{_table_to_markdown(table)}")
        return "\n".join(p for p in parts if p.strip())


MAX_FILE_SIZE = _w.pipeline.max_file_size_mb * 1024 * 1024  # default 200MB


def _table_to_markdown(data: list[list[str]], max_rows: int = 50) -> str:
    """Convert table data to markdown format."""
    if not data:
        return ""
    lines = []
    header = data[0]
    lines.append("| " + " | ".join(header) + " |")
    lines.append("| " + " | ".join(["---"] * len(header)) + " |")
    for row in data[1:max_rows]:
        padded_row = row + [""] * (len(header) - len(row))
        padded_row = padded_row[: len(header)]
        lines.append("| " + " | ".join(padded_row) + " |")
    if len(data) > max_rows:
        lines.append(f"... ({len(data) - max_rows} rows omitted)")
    return "\n".join(lines)


def _extract_table_data(table) -> list[list[str]]:
    """Extract table rows as list of lists."""
    return [[cell.text.strip() for cell in row.cells] for row in table.rows]


def _parse_text(data: bytes) -> str:
    """Parse text files with UTF-8 / EUC-KR fallback."""
    try:
        return data.decode("utf-8")
    except UnicodeDecodeError:
        return data.decode("euc-kr", errors="ignore")


def _parse_xlsx(data: bytes, _filename: str) -> str:
    """Parse XLSX using openpyxl."""
    import openpyxl

    wb = openpyxl.load_workbook(io.BytesIO(data), data_only=True)
    texts = []

    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        texts.append(f"[Sheet: {sheet_name}]")
        sheet_data = []
        for row in sheet.iter_rows(values_only=True):
            row_data = [str(cell) if cell is not None else "" for cell in row]
            if any(cell.strip() for cell in row_data):
                sheet_data.append(row_data)
        if sheet_data:
            texts.append(_table_to_markdown(sheet_data))

    wb.close()
    return "\n\n".join(texts)


def _format_docx_paragraph(para) -> str | None:
    """Format a DOCX paragraph, returning None if empty."""
    if not para.text.strip():
        return None
    if para.style and para.style.name.startswith("Heading"):
        level = para.style.name[-1] if para.style.name[-1].isdigit() else "1"
        return f"{'#' * int(level)} {para.text}"
    return para.text


def _parse_docx(data: bytes, _filename: str) -> str:
    """Parse DOCX using python-docx."""
    from docx import Document

    try:
        doc = Document(io.BytesIO(data))
    except Exception as e:
        raise ValueError(f"DOCX open failed (corrupt?): {e}") from e

    texts = [t for p in doc.paragraphs if (t := _format_docx_paragraph(p))]

    for table in doc.tables:
        table_data = _extract_table_data(table)
        if table_data:
            texts.append(_table_to_markdown(table_data))

    return "\n\n".join(texts)


def _convert_ppt_to_pptx(data: bytes, filename: str) -> bytes | None:
    """Convert legacy .ppt to .pptx using LibreOffice CLI.

    Returns the .pptx file bytes on success, or None if conversion fails
    (e.g. LibreOffice not installed).
    """
    soffice = shutil.which("soffice")
    if soffice is None:
        logger.warning(
            "LibreOffice (soffice) not found on PATH. Cannot convert .ppt file: %s. "
            "Install LibreOffice to enable .ppt support.",
            filename,
        )
        return None

    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            ppt_path = Path(tmpdir) / Path(filename).name
            ppt_path.write_bytes(data)

            result = subprocess.run(
                [
                    soffice,
                    "--headless",
                    "--convert-to",
                    "pptx",
                    "--outdir",
                    tmpdir,
                    str(ppt_path),
                ],
                capture_output=True,
                timeout=_w.timeouts.httpx_default,
            )

            if result.returncode != 0:
                stderr = result.stderr.decode("utf-8", errors="replace").strip()
                logger.warning(
                    "LibreOffice conversion failed for %s (exit %d): %s",
                    filename,
                    result.returncode,
                    stderr,
                )
                return None

            pptx_path = ppt_path.with_suffix(_EXT_PPTX)
            if not pptx_path.exists():
                logger.warning(
                    "LibreOffice conversion produced no output for %s", filename
                )
                return None

            converted_bytes = pptx_path.read_bytes()
            logger.info(
                "Converted .ppt to .pptx: %s (%d bytes)", filename, len(converted_bytes)
            )
            return converted_bytes

    except subprocess.TimeoutExpired:
        logger.warning(
            "LibreOffice conversion timed out (30s) for %s", filename
        )
        return None
    except (RuntimeError, OSError, ValueError, TypeError, KeyError, AttributeError) as e:
        logger.warning("LibreOffice conversion error for %s: %s", filename, e)
        return None


# ---------------------------------------------------------------------------
# OCR / CV Pipeline routing
# ---------------------------------------------------------------------------


def _resize_image(raw: bytes, scale: float = 0.75) -> bytes | None:
    """Resize image as fallback when OCR fails on the original."""
    try:
        from PIL import Image
        img = Image.open(io.BytesIO(raw))
        new_w = max(int(img.width * scale), 32)
        new_h = max(int(img.height * scale), 32)
        img = img.resize((new_w, new_h), Image.LANCZOS)
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        return buf.getvalue()
    except (RuntimeError, OSError, ValueError, TypeError, KeyError, AttributeError):
        return None


def _ocr_request(client, endpoint: str, payload: bytes) -> dict | None:
    """Send OCR request. Returns parsed JSON or None on failure."""
    import base64
    b64_image = base64.b64encode(payload).decode("utf-8")
    try:
        resp = client.post(endpoint, json={"image": b64_image}, timeout=_w.timeouts.httpx_ocr)
        resp.raise_for_status()
        return resp.json()
    except (RuntimeError, OSError, ValueError, TypeError, KeyError, AttributeError):
        return None


def _prepare_image_bytes(img_bytes: bytes, index: int) -> bytes | None:
    """Validate, decode, and convert image to PNG. Returns None if invalid."""
    from PIL import Image
    _MIN_W, _MIN_H = 20, 20

    try:
        _img = Image.open(io.BytesIO(img_bytes))
    except (RuntimeError, OSError, ValueError, TypeError, KeyError, AttributeError):
        logger.debug("Skipping image %d: cannot decode", index)
        return None
    if _img.width < _MIN_W or _img.height < _MIN_H:
        logger.debug("Skipping image %d: too small (%dx%d)", index, _img.width, _img.height)
        return None
    if _img.mode not in ("RGB", "L"):
        _img = _img.convert("RGB")
    _buf = io.BytesIO()
    _img.save(_buf, format="PNG")
    return _buf.getvalue()


def _extract_text_from_boxes(boxes: list[dict], min_conf: float, index: int) -> str:
    """Extract text from OCR boxes with confidence filtering."""
    filtered_texts = []
    dropped = 0
    for box in boxes:
        conf = box.get("confidence", 0.0)
        box_text = box.get("text", "")
        if conf >= min_conf and box_text.strip():
            filtered_texts.append(box_text)
        elif box_text.strip():
            dropped += 1
    if dropped > 0:
        logger.info(
            "OCR confidence filter: kept %d, dropped %d (< %.0f%%) for image %d",
            len(filtered_texts), dropped, min_conf * 100, index,
        )
    return " ".join(filtered_texts)


def _extract_text_fallback(result: dict) -> str:
    """Extract text from OCR result when no boxes are available."""
    text_lines = result.get("texts", result.get("result", []))
    if isinstance(text_lines, list):
        return " ".join(str(t) for t in text_lines if t)
    if isinstance(text_lines, str):
        return text_lines
    return str(text_lines)


def _extract_vision_analysis(
    result: dict, index: int,
    ocr_texts: list[str], visual_analyses: list[dict[str, Any]],
) -> None:
    """Extract vision analysis (shapes/arrows/mappings) from OCR result."""
    shapes = result.get("shapes", [])
    arrows = result.get("arrows", [])
    mappings = result.get("text_shape_mappings", [])
    if not shapes and not arrows and not mappings:
        return

    visual_analyses.append({
        "image_index": index,
        "shapes": shapes,
        "arrows": arrows,
        "text_shape_mappings": mappings,
        "shape_count": result.get("shape_count", len(shapes)),
        "arrow_count": result.get("arrow_count", len(arrows)),
        "ocr_confidence": result.get("ocr_confidence", 0.0),
    })

    desc_parts: list[str] = []
    if shapes:
        shape_types = [s.get("type", "unknown") for s in shapes]
        desc_parts.append(f"Shapes: {', '.join(shape_types)}")
    if arrows:
        desc_parts.append(f"Arrows: {len(arrows)} connections")
    if mappings:
        for m in mappings:
            shape_label = m.get("shape_type", "shape")
            texts_in = m.get("texts", m.get("text", ""))
            if texts_in:
                desc_parts.append(f"  [{shape_label}] {texts_in}")
    if desc_parts:
        diagram_desc = "\n".join(desc_parts)
        ocr_texts.append(f"[Image {index} Diagram]\n{diagram_desc}")


def _process_single_image_ocr(
    client,
    endpoint: str,
    img_bytes: bytes,
    index: int,
    min_conf: float,
    vision_enabled: bool,
    ocr_texts: list[str],
    visual_analyses: list[dict[str, Any]],
) -> None:
    """Process a single image through OCR with retry on failure."""
    try:
        prepared = _prepare_image_bytes(img_bytes, index)
        if prepared is None:
            return

        result = _ocr_request(client, endpoint, prepared)
        if result is None:
            resized = _resize_image(prepared)
            if resized:
                logger.info("Retrying image %d with resized version", index)
                result = _ocr_request(client, endpoint, resized)

        if result is None:
            logger.warning("PaddleOCR failed for image %d after retry, skipping", index)
            return

        boxes = result.get("boxes", [])
        text = (
            _extract_text_from_boxes(boxes, min_conf, index)
            if boxes
            else _extract_text_fallback(result)
        )

        if text.strip():
            ocr_texts.append(f"[Image {index} OCR] {text}")
            logger.info("OCR success for image %d: %d chars", index, len(text))

        if vision_enabled:
            _extract_vision_analysis(result, index, ocr_texts, visual_analyses)

    except (RuntimeError, OSError, ValueError, TypeError, KeyError, AttributeError) as e:
        logger.warning("PaddleOCR API failed for image %d: %s", index, e)


def _process_images_ocr(
    images: list[bytes],
) -> tuple[str, list[dict[str, Any]]]:
    """Route extracted images through PaddleOCR Docker API server."""
    import httpx
    import os

    if not images:
        return "", []

    ocr_texts: list[str] = []
    visual_analyses: list[dict[str, Any]] = []
    base_url = os.getenv("PADDLEOCR_API_URL", "http://localhost:8866")
    if base_url.endswith("/ocr") or base_url.endswith("/analyze"):
        base_url = base_url.rsplit("/", 1)[0]

    vision_enabled = _w.ocr.enable_vision_analysis
    endpoint = f"{base_url}/analyze" if vision_enabled else f"{base_url}/ocr"
    min_conf = float(os.getenv("OCR_MIN_CONFIDENCE", "0.65"))
    client = httpx.Client(timeout=_w.timeouts.httpx_ocr)

    for i, img_bytes in enumerate(images):
        _process_single_image_ocr(
            client, endpoint, img_bytes, i + 1, min_conf, vision_enabled,
            ocr_texts, visual_analyses,
        )

    client.close()
    return "\n".join(ocr_texts), visual_analyses
